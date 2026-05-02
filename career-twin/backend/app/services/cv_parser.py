import fitz  # PyMuPDF
import pdfplumber
import io


class UnsupportedFileTypeError(ValueError):
    pass


def extract_text_from_pdf_bytes(data: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF. Falls back to pdfplumber."""
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
        if text.strip():
            return text.strip()
    except Exception:
        pass

    # Fallback: pdfplumber
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        text = "\n".join(
            page.extract_text() or "" for page in pdf.pages
        )
    return text.strip()


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    """Route extraction by file extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf_bytes(data)
    if lower.endswith(".docx"):
        return _extract_text_from_docx(data)
    if lower.endswith(".doc"):
        raise UnsupportedFileTypeError(
            "Legacy .doc format is not supported. Please convert it to .docx and try again."
        )
    raise UnsupportedFileTypeError(f"Unsupported file type: {filename}")


def _extract_text_from_docx(data: bytes) -> str:
    """Extract text from DOCX bytes using python-docx, reading paragraphs and tables."""
    try:
        import docx
    except ImportError:
        raise UnsupportedFileTypeError(
            "python-docx is not installed. Please add 'python-docx==1.1.2' to your requirements."
        )

    try:
        doc = docx.Document(io.BytesIO(data))
        full_text = []

        # 1. 提取普通段落
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # 2. 提取表格（简历常用表格排版）
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        # 避免重复提取单元格内嵌套段落的内容
                        row_text.append(cell.text.strip())
                if row_text:
                    # 使用管道符拼接单元格，帮助大模型更好地理解对应关系
                    full_text.append(" | ".join(dict.fromkeys(row_text)))

        return "\n".join(full_text).strip()

    except Exception as e:
        raise UnsupportedFileTypeError(
            f"Failed to parse the Word file. It might be corrupted or in an unsupported format. Error: {str(e)}"
        )